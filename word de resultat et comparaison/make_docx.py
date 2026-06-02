from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r'C:\_workingDir\_SF\test flexion'
IMG_KRG = os.path.join(BASE, 'etat_limite_KRG.png')
IMG_GEK = os.path.join(BASE, 'etat_limite_GEK.png')
OUT = os.path.join(BASE, 'comparaison_KRG_GEK.docx')

doc = Document()

# ---- marges ----
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def bold_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p

def normal(doc, text):
    doc.add_paragraph(text)

def shade_row(row, hex_color='D9E1F2'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

# ---- Titre ----
heading(doc, 'Comparaison KRG pur vs GEK pur — DOE fixé n0=15, F=0.210 MN', level=1)

p = doc.add_paragraph()
p.add_run('Date : ').bold = True
p.add_run('23 avril 2026')
p = doc.add_paragraph()
p.add_run('Configuration commune : ').bold = True
p.add_run('n0=15, DOE fixé (U-space), F=0.210 MN, solver AbdoRackwitz')
p = doc.add_paragraph()
p.add_run('Référence HF : ').bold = True
p.add_run('β=3.784, Pf=7.73e-05, u*=[-0.526, -3.747], n_iter=21')

doc.add_paragraph()

# ---- DOE ----
heading(doc, 'DOE fixé utilisé (n0=15, U-space)', level=2)

doe_lines = [
    '[ 1.0272625484832025,  0.3251235065050853]',
    '[ 0.2588934150948534, -1.6856336900013655]',
    '[-0.7900915845657982,  1.8047217395005692]',
    '[-0.0301755082064849,  1.3223984111477798]',
    '[-1.8073810055112547, -1.1012751718677385]',
    '[-0.2377471223963969, -0.4914312425631510]',
    '[ 0.7216266145109314,  1.0830320538875535]',
    '[ 0.4776729449462016, -0.2656508781535193]',
    '[-0.8730465106774573,  0.6497494474356423]',
    '[-1.1677174906609287,  0.0310652111349381]',
    '[ 1.1194425579629474, -0.7943643305093363]',
    '[ 0.1857520921586401,  0.4724170659386679]',
    '[-0.5669380193636159, -1.4858232340964800]',
    '[ 2.9454553139272623, -0.1582987245612891]',
    '[-0.2947626989079067,  0.1355018527305618]',
]
code_text = 'U_doe_fixed = ot.Sample([\n' + '\n'.join(f'    {l},' for l in doe_lines) + '\n])'
p = doc.add_paragraph(code_text)
p.style = doc.styles['No Spacing']
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_paragraph()

# ---- Partie 1 ----
heading(doc, 'Partie 1 — Résultats détaillés', level=2)

rows_p1 = [
    ('Paramètre', 'HF (référence)', 'KRG pur', 'GEK pur'),
    ('n points DOE', '—', '15 (fixé)', '15 (fixé)'),
    ('n iter FORM', '21', '24', '25'),
    ('n_appels HF (FORM)', '—', '0', '0'),
    ('fc* (MPa)', '—', '26.7006', '30.9316'),
    ('fy* (MPa)', '—', '474.9796', '548.7802'),
    ('u*', '[-0.526, -3.747]', '[-1.767, -4.755]', '[-0.372, -2.307]'),
    ('dg/du_fc en u*', '—', '0.004850', '0.006877'),
    ('dg/du_fy en u*', '—', '0.002416', '0.041592'),
    ('Importance fc', '—', '12.13%', '2.53%'),
    ('Importance fy', '—', '87.87%', '97.47%'),
    ('β (FORM)', '3.784', '5.072', '2.337'),
    ('Pf (FORM)', '7.73e-05', '1.96e-07', '9.73e-03'),
    ('g_meta(u*)', '~0', '+8.7e-05', '+0.0606'),
    ('g_HF(u*)', '~0', '-4.98e-02', '+0.0608'),
    ('Erreur relative g', '—', '100.2%', '0.32%'),
    ('u* FOSM (depuis u=0)', '—', '[-0.643, -3.728]', '[-0.643, -3.728]'),
    ('Erreur FOSM', '—', '30.00%', '61.93%'),
    ('Ecart β vs HF', '0%', '+1.288 (+34.0%)', '-1.447 (-38.2%)'),
    ('do_warm_start', '—', 'False', 'False'),
]

t1 = doc.add_table(rows=len(rows_p1), cols=4)
t1.style = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(rows_p1):
    row = t1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
    if i == 0:
        shade_row(row, 'BDD7EE')
    elif i % 2 == 0:
        shade_row(row, 'F2F2F2')

doc.add_paragraph()

# ---- Partie 2 ----
heading(doc, 'Partie 2 — Bilan de comparaison', level=2)

rows_p2 = [
    ('', 'HF', 'KRG pur', 'GEK pur'),
    ('β', '3.784', '5.072', '2.337'),
    ('Erreur β', '0%', '+34.0%', '-38.2%'),
    ('u*', '[-0.53, -3.75]', '[-1.77, -4.75]', '[-0.37, -2.31]'),
    ('g_meta(u*)', '~0', '+8.7e-05', '+0.0606'),
    ('g_HF(u*)', '~0', '-4.98e-02', '+0.0608'),
    ('Erreur locale g', '—', '100.2%', '0.32%'),
]

t2 = doc.add_table(rows=len(rows_p2), cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(rows_p2):
    row = t2.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
    if i == 0:
        shade_row(row, 'BDD7EE')
    elif i % 2 == 0:
        shade_row(row, 'F2F2F2')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Conclusion : ').bold = True
p.add_run(
    'Sur ce DOE fixé n0=15, KRG pur surestime β (+34%) et GEK pur le sous-estime (-38%). '
    'KRG converge sur sa surface g=0 (g_meta=8.7e-05) mais la surface est mal placée (g_HF=-0.050, erreur 100%). '
    'GEK pur ne converge pas non plus sur sa surface (g_meta=+0.061) — FORM s\'est arrêté à un point non-défaillant. '
    'Paradoxalement, GEK est localement précis en u* (erreur 0.32%) mais le métamodèle est mal conditionné '
    'globalement avec n0=15. L\'activation du warm start pour GEK et l\'ajout de GEPCK (GEK+PCE) sont les prochaines étapes.'
)

doc.add_paragraph()

# ---- Partie 3 — Visualisations ----
heading(doc, 'Partie 3 — Surfaces limites (g_HF=0 et g_GP=0)', level=2)

# Table 1 ligne × 2 colonnes pour images côte à côte
t3 = doc.add_table(rows=2, cols=2)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

# Ligne 0 : titres
t3.rows[0].cells[0].text = 'KRG pur'
t3.rows[0].cells[1].text = 'GEK pur'
for cell in t3.rows[0].cells:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
shade_row(t3.rows[0], 'BDD7EE')

# Ligne 1 : images
img_width = Cm(8.5)
cell_krg = t3.rows[1].cells[0]
cell_gek = t3.rows[1].cells[1]

para_krg = cell_krg.paragraphs[0]
para_krg.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_krg = para_krg.add_run()
run_krg.add_picture(IMG_KRG, width=img_width)

para_gek = cell_gek.paragraphs[0]
para_gek.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_gek = para_gek.add_run()
run_gek.add_picture(IMG_GEK, width=img_width)

doc.save(OUT)
print(f'Document sauvegardé : {OUT}')
